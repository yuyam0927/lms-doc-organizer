from pathlib import Path

from docx import Document

from lms_document_to_md_parser.parsers.docx_parser import docx_to_markdown


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("見出し1", level=1)
    doc.add_paragraph("本文段落")
    doc.add_paragraph("箇条書き項目", style="List Bullet")
    doc.add_paragraph("")  # empty paragraph should be dropped
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "h1"
    table.rows[0].cells[1].text = "h2"
    table.rows[1].cells[0].text = "a|b"
    table.rows[1].cells[1].text = "line1\nline2"
    doc.save(str(path))


def test_docx_to_markdown_headings_paragraphs_lists_and_tables(tmp_path):
    docx_path = tmp_path / "sample.docx"
    _make_docx(docx_path)

    md = docx_to_markdown(docx_path)

    assert "# 見出し1" in md
    assert "本文段落" in md
    assert "- 箇条書き項目" in md
    assert "| h1 | h2 |" in md
    assert "| a\\|b | line1<br>line2 |" in md


def test_docx_to_markdown_skips_empty_paragraphs(tmp_path):
    docx_path = tmp_path / "empty.docx"
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.save(str(docx_path))

    md = docx_to_markdown(docx_path)

    assert md.strip() == ""
